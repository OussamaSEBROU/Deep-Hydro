import streamlit as st
import pandas as pd
import numpy as np
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_squared_error, mean_absolute_error, mean_absolute_percentage_error
import tensorflow as tf
from tensorflow.keras.models import Sequential, load_model
from tensorflow.keras.layers import LSTM, Dropout, Dense
from tensorflow.keras.callbacks import EarlyStopping
import plotly.graph_objects as go
import google.generativeai as genai
import io
import base64
import time
import os
import json
import datetime
import uuid
import firebase_admin
from firebase_admin import credentials, db
import requests
import plotly.express as px
from dotenv import load_dotenv
import hashlib
import streamlit.components.v1 as components 
import matplotlib.pyplot as plt 
from statsmodels.graphics.tsaplots import plot_acf 
from docx import Document 
from docx.shared import Inches, Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH 

# --- Constants ---
# ADVANCED_FEATURE_LIMIT removed as features are now unlimited.

# --- Firebase Configuration --- 
def initialize_firebase():
    """
    Initialize Firebase with secure credential management.
    Loads credentials from environment variables for secure deployment.
    """
    load_dotenv()
    if not firebase_admin._apps:
        try:
            firebase_creds_json = os.getenv("FIREBASE_SERVICE_ACCOUNT")
            if firebase_creds_json:
                cred_dict = json.loads(firebase_creds_json)
                cred = credentials.Certificate(cred_dict)
                firebase_url = os.getenv("FIREBASE_DATABASE_URL", f"https://{cred_dict.get('project_id')}-default-rtdb.firebaseio.com/")
                firebase_admin.initialize_app(cred, {
                    "databaseURL": firebase_url
                })
                # st.success("Firebase initialized successfully.") # Optional: for debugging
                return True
            else:
                st.warning("Firebase credentials not found. Analytics and usage tracking are disabled.")
                return False
        except Exception as e:
            st.warning(f"Firebase initialization error: {e}. Analytics and usage tracking are disabled.")
            return False
    return True

# --- User Identification & Tracking --- 
def get_client_ip():
    """Get the client's IP address if available."""
    try:
        # Use a reliable service to get IP
        response = requests.get('https://api.ipify.org?format=json', timeout=3)
        if response.status_code == 200:
            return response.json().get('ip', 'Unknown')
        return "Unknown"
    except Exception:
        return "Unknown"

def get_persistent_user_id():
    """Generate or retrieve a persistent user ID. Always generates an anonymous ID."""
    if 'persistent_user_id' not in st.session_state:
        ip_address = get_client_ip()
        user_agent = st.session_state.get('user_agent', 'Unknown')
        
        # Create a stable hash from IP and user agent
        hash_input = f"{ip_address}-{user_agent}"
        hashed_id = hashlib.sha256(hash_input.encode()).hexdigest()
        st.session_state.persistent_user_id = f"anon_{hashed_id}"
        
    return st.session_state.persistent_user_id

def get_or_create_user_profile(user_id):
    """Get user profile from Firebase or create a new one."""
    if not firebase_admin._apps:
        return None, False # Indicate Firebase not available
    
    try:
        ref = db.reference(f'users/{user_id}')
        profile = ref.get()
        is_new_user = False
        if profile is None:
            is_new_user = True
            profile = {
                'user_id': user_id,
                'first_visit': datetime.datetime.now().isoformat(),
                'visit_count': 1,
                'feature_usage_count': 0, # Still track, but not for limits
                'last_visit': datetime.datetime.now().isoformat(),
            }
            ref.set(profile)
        else:
            # Update visit count and last visit time if it's a new session
            if 'session_visit_logged' not in st.session_state:
                profile['visit_count'] = profile.get('visit_count', 0) + 1
                profile['last_visit'] = datetime.datetime.now().isoformat()
                ref.update({'visit_count': profile['visit_count'], 
                            'last_visit': profile['last_visit']})
                st.session_state.session_visit_logged = True # Mark visit as logged for this session
            
        return profile, is_new_user
    except Exception as e:
        st.warning(f"Firebase error getting/creating user profile for {user_id}: {e}")
        return None, False

# Removed increment_feature_usage and check_feature_access as per requirements (unlimited features)
# Removed show_google_login_button and all related Google Auth logic.

# --- Visitor Analytics Functions --- 
def get_session_id():
    """Create or retrieve a unique session ID."""
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    return st.session_id

def log_visitor_activity(page_name, action="page_view", feature_used=None):
    """Log visitor activity to Firebase."""
    if not firebase_admin._apps:
        return
    try:
        user_id = get_persistent_user_id()
        profile, _ = get_or_create_user_profile(user_id)
        
        ref = db.reference('visitors_log') 
        log_id = str(uuid.uuid4())
        timestamp = datetime.datetime.now().isoformat()
        session_id = get_session_id()
        user_agent = st.session_state.get('user_agent', 'Unknown')
        ip_address = get_client_ip() 

        is_authenticated = False 

        log_data = {
            'timestamp': timestamp,
            'persistent_user_id': user_id,
            'is_authenticated': is_authenticated,
            'visit_count': profile.get('visit_count', 1) if profile else 1,
            'ip_address': ip_address,
            'page': page_name,
            'action': action,
            'feature_used': feature_used,
            'session_id': session_id,
            'user_agent': user_agent,
        }
        ref.child(log_id).set(log_data)
    except Exception as e:
        pass # Silently fail logging

def fetch_visitor_logs():
    """Fetch visitor logs from Firebase."""
    if not firebase_admin._apps:
        return pd.DataFrame()
    try:
        ref = db.reference('visitors_log')
        visitors_data = ref.get()
        if not visitors_data: return pd.DataFrame()
        visitors_list = [dict(log_id=log_id, **data) for log_id, data in visitors_data.items()]
        df = pd.DataFrame(visitors_list)
        df['timestamp'] = pd.to_datetime(df['timestamp'])
        df = df.sort_values('timestamp', ascending=False)
        return df
    except Exception as e:
        st.error(f"Error fetching visitor logs: {e}")
        return pd.DataFrame()

def create_visitor_charts(visitor_df):
    """Create visualizations of visitor data."""
    if visitor_df.empty: return []
    figures = []
    try:
        df = visitor_df.copy()
        df['date'] = df['timestamp'].dt.date
        
        # Daily visitors
        daily_visitors = df.groupby('date')['persistent_user_id'].nunique().reset_index(name='unique_users')
        daily_visitors['date'] = pd.to_datetime(daily_visitors['date'])
        fig1 = px.line(daily_visitors, x='date', y='unique_users', title='Daily Unique Visitors', labels={'unique_users': 'Unique Users', 'date': 'Date'})
        figures.append(fig1)
        
        # Action counts
        action_counts = df['action'].value_counts().reset_index()
        action_counts.columns = ['action', 'count']
        fig2 = px.bar(action_counts, x='action', y='count', title='Activity Counts by Action', labels={'count': 'Number of Times', 'action': 'Action Type'})
        figures.append(fig2)

        # Auth status - simplified as authentication is removed
        # Removed the 'Auth status' pie chart since Google Sign-in is no longer used.

        # Hourly activity heatmap
        try:
            df['hour'] = df['timestamp'].dt.hour
            df['day_of_week'] = df['timestamp'].dt.day_name()
            day_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            hourly_activity = df.groupby(['day_of_week', 'hour']).size().reset_index(name='count')
            
            if len(hourly_activity) > 0:
                hourly_pivot = hourly_activity.pivot_table(values='count', index='day_of_week', columns='hour', fill_value=0)
                available_days = set(hourly_pivot.index) & set(day_order)
                ordered_available_days = [day for day in day_order if day in available_days]
                hourly_pivot = hourly_pivot.reindex(ordered_available_days)
                available_hours = sorted(hourly_pivot.columns)
                fig4 = px.imshow(hourly_pivot, labels=dict(x="Hour of Day", y="Day of Week", color="Activity Count"),
                                x=[str(h) for h in available_hours], y=ordered_available_days, title="Visitor Activity by Hour and Day")
                figures.append(fig4)
            else:
                fig4 = go.Figure().update_layout(title="Visitor Activity by Hour and Day (No Data)")
                fig4.add_annotation(text="Not enough data", xref="paper", yref="paper", x=0.5, y=0.5, showarrow=False)
                figures.append(fig4)
        except Exception as heatmap_err:
            st.warning(f"Could not generate hourly activity heatmap: {heatmap_err}")
            fig4 = go.Figure().update_layout(title="Visitor Activity by Hour and Day (Error)")
            fig4.add_annotation(text="Error generating heatmap", xref="paper", yref="paper", x=0.5, y=0.5, showarrow=False)
            figures.append(fig4)
            
    except Exception as e:
        st.error(f"Error creating visitor charts: {e}")
        return []
    return figures

# --- Admin Analytics Dashboard --- 
def render_admin_analytics():
    """Render the admin analytics dashboard."""
    st.header("Admin Analytics Dashboard")
    if 'admin_authenticated' not in st.session_state: st.session_state.admin_authenticated = False
    
    if not st.session_state.admin_authenticated:
        st.info("Admin access required.")
        admin_password = st.text_input("Admin Password", type="password", key="admin_pass_input")
        # Added a login button explicitly
        if st.button("Login", key="admin_login_btn", use_container_width=True): 
            correct_password = os.getenv("ADMIN_PASSWORD", "admin123")
            if admin_password == correct_password:
                st.session_state.admin_authenticated = True
                st.rerun()
            else: st.error("Invalid password")
    else:
        visitor_df = fetch_visitor_logs()
        if visitor_df.empty: st.info("No visitor data available yet."); return
        
        st.subheader("Visitor Statistics")
        col1, col2, col3 = st.columns(3)
        with col1: st.metric("Total Activities Logged", len(visitor_df))
        with col2: st.metric("Unique Visitors", visitor_df['persistent_user_id'].nunique())
        with col3:
            if 'date' not in visitor_df.columns: visitor_df['date'] = visitor_df['timestamp'].dt.date
            today = datetime.datetime.now().date()
            today_visitors = visitor_df[visitor_df['date'] == today]['persistent_user_id'].nunique()
            st.metric("Today's Unique Visitors", today_visitors)
        
        st.subheader("Visitor Analytics")
        try:
            charts = create_visitor_charts(visitor_df)
            for fig in charts: st.plotly_chart(fig, use_container_width=True)
        except Exception as chart_err: st.error(f"Error displaying charts: {chart_err}")
        
        st.subheader("Raw Visitor Data")
        col1_filter, col2_filter = st.columns(2)
        with col1_filter:
            try:
                min_date, max_date = visitor_df['timestamp'].min().date(), visitor_df['timestamp'].max().date()
                date_range = st.date_input("Date Range", [min_date, max_date], min_value=min_date, max_value=max_date, key="admin_date_filter")
            except Exception: date_range = None
        with col2_filter:
            try:
                user_id_options = ['All'] + visitor_df['persistent_user_id'].unique().tolist()
                user_id_filter = st.selectbox("Filter by User ID", options=user_id_options, index=0, key="admin_user_filter")
            except Exception: user_id_filter = 'All'
        
        try:
            filtered_df = visitor_df.copy()
            if date_range and len(date_range) == 2:
                start_date, end_date = date_range
                filtered_df = filtered_df[(filtered_df['timestamp'].dt.date >= start_date) & (filtered_df['timestamp'].dt.date <= end_date)]
            if user_id_filter != 'All':
                filtered_df = filtered_df[filtered_df['persistent_user_id'] == user_id_filter]
            
            # Removed 'is_authenticated' and 'google_email' from display_cols
            display_cols = ['timestamp', 'persistent_user_id', 'visit_count', 'page', 'action', 'feature_used', 'ip_address', 'session_id']
            st.dataframe(filtered_df[[col for col in display_cols if col in filtered_df.columns]])
            
            if st.button("Export Filtered to CSV", key="admin_export_btn"):
                csv = filtered_df[[col for col in display_cols if col in filtered_df.columns]].to_csv(index=False)
                b64 = base64.b64encode(csv.encode()).decode()
                href = f'<a href="data:file/csv;base64,{b64}" download="filtered_visitor_logs.csv">Download CSV File</a>'
                st.markdown(href, unsafe_allow_html=True)
        except Exception as filter_err:
            st.error(f"Error applying filters or displaying data: {filter_err}")
            st.dataframe(visitor_df[['timestamp', 'persistent_user_id', 'action']])

# --- UI Assets --- 
def get_custom_css():
    """Returns the custom CSS string for the application."""
    return """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;500;600;700&display=swap');
    
    html, body, [data-testid="stAppViewContainer"] {
        font-family: 'Montserrat', sans-serif;
    }

    /* Sidebar style preservation */
    .sidebar .block-container {
        font-size: 0.9rem;
    }
    .sidebar h1 {
        font-size: 1.4rem;
    }
    .sidebar h2 {
        font-size: 1.2rem;
    }
    /* ALL sidebar buttons styling - RED */
    /* This rule will target all st.button elements within the sidebar */
    .stSidebar .stButton > button {
        background-color: #FF4B4B !important; /* Red color, !important to override */
        color: white !important;
        border: 1px solid #FF4B4B !important;
        border-radius: 4px;
        font-weight: 500;
        transition: all 0.3s;
        padding: 0.5rem 1rem;
    }
    .stSidebar .stButton > button:hover {
        background-color: #e03e3e !important; /* Darker red on hover */
        border-color: #e03e3e !important;
        box-shadow: 0 2px 5px rgba(255,75,75,0.4) !important; /* More pronounced shadow */
        transform: translateY(-1px) !important; /* Slight lift */
    }

    .sidebar .stFileUploader, .sidebar .stSelectbox, .sidebar .stNumberInput, .sidebar .stCheckbox {
        margin-bottom: 0.8rem; /* Consistent spacing */
    }
    .sidebar .stNumberInput input {
        font-size: 0.9rem; /* Slightly smaller input text */
    }
    .sidebar .stExpander {
        border: none;
        box-shadow: none;
        background-color: transparent;
    }
    .sidebar .stExpander header {
        padding: 0.5rem 0;
        font-weight: 500;
    }
    .sidebar .stExpander div[data-testid="stExpanderDetails"] {
        padding-left: 0.5rem;
    }
    .about-us-header {
        cursor: pointer;
        padding: 0.5rem;
        border-radius: 4px;
        margin-top: 1rem;
        font-weight: 500;
    }
    .about-us-content {
        padding: 0.8rem;
        border-radius: 4px;
        margin-top: 0.5rem;
        font-size: 0.9rem;
    }
    
    /* Main content styles - Professional and compact */
    .main .block-container { 
        padding-top: 1.5rem; /* Increased top padding */
        padding-bottom: 1.5rem; /* Increased bottom padding */
        padding-left: 2rem; /* Consistent side padding */
        padding-right: 2rem; /* Consistent side padding */
    }
    h1 { 
        font-weight: 700; 
        font-size: 2.2rem; /* Larger and bolder for main titles */
        color: var(--primary-color); /* Main titles in blue */
        margin-bottom: 1.5rem; /* Space below title */
    }
    h2 { 
        font-weight: 600; 
        font-size: 1.8rem; 
        color: var(--text-color); 
        margin-top: 2rem; /* More space above subheaders */
        margin-bottom: 1rem;
    }
    h3 { 
        font-weight: 600; 
        font-size: 1.4rem; 
        color: var(--text-color); 
        margin-top: 1.5rem; /* More space above h3 */
        margin-bottom: 0.8rem;
    }
    h4 { 
        font-weight: 500; 
        font-size: 1.2rem; 
        color: var(--text-color); 
        margin-top: 1rem;
        margin-bottom: 0.5rem;
    }

    /* Primary buttons outside sidebar - blue */
    /* This rule ensures buttons in the main area remain blue unless specifically overridden */
    .stButton:not(.stSidebar .stButton) > button { 
        border-radius: 8px; /* Slightly more rounded buttons */
        font-weight: 600; /* Bolder button text */
        transition: all 0.3s ease-in-out; 
        padding: 0.7rem 1.5rem; /* Larger padding for buttons */
        border: 1px solid var(--primary-color); /* Outline button for professional look */
        background-color: var(--primary-color);
        color: white; /* White text on primary background */
    }
    .stButton:not(.stSidebar .stButton) > button:hover { 
        opacity: 0.9; 
        box-shadow: 0 4px 12px rgba(0,0,0,0.1); /* More pronounced shadow on hover */
        transform: translateY(-2px); /* Slight lift effect */
    }

    /* Style for download buttons (e.g., CSV/XLSX) if different styling is desired) */
    .stDownloadButton > button {
        background-color: var(--background-color);
        color: var(--primary-color);
        border: 1px solid var(--primary-color);
    }
    .stDownloadButton > button:hover {
        background-color: var(--primary-color);
        color: white;
    }

    .css-1d391kg, .css-12oz5g7 { padding: 1rem; } 
    .card-container { 
        border-radius: 12px; /* More rounded cards */
        padding: 1.5rem; /* Increased card padding */
        margin-bottom: 1.5rem; /* More space between cards */
        box-shadow: 0 4px 12px rgba(0,0,0,0.08); /* Stronger, softer shadow */
        background-color: var(--secondary-background-color); /* Use secondary background for cards */
        border: 1px solid rgba(128, 128, 128, 0.1); /* Subtle border */
    }

    .chat-message { 
        padding: 1rem; 
        border-radius: 12px; /* Rounded chat messages */
        margin-bottom: 0.75rem; 
        position: relative; 
        font-size: 0.95rem; /* Slightly larger text */
    }
    
    .user-message { 
        border-left: 5px solid var(--primary-color); /* Thicker border */
        background-color: var(--secondary-background-color); 
        color: var(--text-color);
    }
    .ai-message { 
        border-left: 5px solid #78909C; /* A neutral gray, thicker border */
        background-color: var(--secondary-background-color); 
        color: var(--text-color);
    }

    .chat-message:active { opacity: 0.7; }
    .copy-tooltip { 
        position: absolute; 
        top: 0.5rem; 
        right: 0.5rem; 
        padding: 0.2rem 0.5rem; 
        border-radius: 6px; 
        font-size: 0.75rem; 
        display: none; 
        background-color: #555; 
        color: white; 
        z-index: 10; 
        box-shadow: 0 2px 4px rgba(0,0,0,0.2);
    }
    .chat-message:active .copy-tooltip { display: block; }
    .stTabs [data-baseweb="tab-list"] { 
        gap: 8px; /* More space between tabs */
        margin-bottom: 1.5rem; /* Space below tabs */
    }
    .stTabs [data-baseweb="tab"] { 
        padding: 0.75rem 1.5rem; /* Larger tab clickable area */
        border-radius: 8px 8px 0 0; /* More rounded tabs */
        font-weight: 600; /* Bolder tab text */
        background-color: var(--secondary-background-color); /* Tabs match card background */
        color: var(--text-color);
        transition: all 0.2s ease-in-out;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        border-bottom: 3px solid var(--primary-color); /* Active tab indicator */
        color: var(--primary-color);
        background-color: var(--background-color); /* Active tab slightly different background */
    }
    .stTabs [data-baseweb="tab"]:hover {
        background-color: var(--background-color);
    }

    [data-testid="stMetricValue"] { 
        font-weight: 700; 
        font-size: 1.5rem; /* Larger metric values */
        color: var(--primary-color); /* Metrics in blue */
    }
    [data-testid="stMetricLabel"] {
        font-size: 0.9rem;
        color: var(--text-color);
    }
    
    /* Theme-aware app intro */
    .app-intro { 
        padding: 2rem; /* More spacious intro */
        border-radius: 15px; /* Very rounded intro */
        margin-bottom: 2.5rem; 
        border-left: 8px solid var(--primary-color); /* Thicker blue border */
        background-color: var(--secondary-background-color); 
        color: var(--text-color);
        box-shadow: 0 6px 20px rgba(0,0,0,0.15); /* Stronger intro shadow */
    }
    .app-intro h3 {
        font-size: 1.8rem;
        color: var(--primary-color); /* Intro heading in blue */
        margin-bottom: 1rem;
    }

    .blue-text {
        color: #1E88E5; /* A standard blue */
    }
    
    /* User Info Display - Removed as Google Sign-in is no longer used */
    
    /* General input styling */
    .stFileUploader label, .stSelectbox label, .stNumberInput label, .stCheckbox label, .stTextInput label {
        font-weight: 500; /* Bolder labels */
        color: var(--text-color);
    }
    .stFileUploader div[data-testid="stFileUploadDropzone"],
    .stTextInput input, .stNumberInput input, .stSelectbox div[data-baseweb="select"] {
        border-radius: 8px; /* Rounded inputs */
        border: 1px solid var(--secondary-background-color); /* Subtle border for inputs */
        padding: 0.5rem 1rem;
    }
    .stFileUploader div[data-testid="stFileUploadDropzone"]:hover {
        border-color: var(--primary-color); /* Highlight on hover */
    }
    </style>
    """

def get_custom_javascript():
    """Returns the custom JavaScript string for UI interactions."""
    return """
    <script>
    // Function to copy text to clipboard
    function copyToClipboard(text) {
        const textarea = document.createElement('textarea');
        textarea.value = text;
        document.body.appendChild(textarea);
        textarea.select();
        document.execCommand('copy');
        document.body.removeChild(textarea);
    }
    
    // Add event listeners after DOM is ready
    document.addEventListener('DOMContentLoaded', function() {
        // Use a small delay to ensure Streamlit elements are fully rendered
        setTimeout(function() {
            // Copy functionality for chat messages
            const chatMessages = document.querySelectorAll('.chat-message');
            chatMessages.forEach(function(message) {
                // Add tooltip element dynamically if not present
                if (!message.querySelector('.copy-tooltip')) {
                    const tooltip = document.createElement('span');
                    tooltip.className = 'copy-tooltip';
                    tooltip.textContent = 'Copied!';
                    message.appendChild(tooltip);
                }
                
                let longPressTimer;
                // Touch events for mobile long press
                message.addEventListener('touchstart', function(e) {
                    longPressTimer = setTimeout(() => {
                        const textToCopy = this.innerText.replace('Copied!', '').trim();
                        copyToClipboard(textToCopy);
                        const tooltip = this.querySelector('.copy-tooltip');
                        if (tooltip) {
                            tooltip.style.display = 'block';
                            setTimeout(() => { tooltip.style.display = 'none'; }, 1500);
                        }
                    }, 500); // 500ms threshold
                });
                message.addEventListener('touchend', function() { clearTimeout(longPressTimer); });
                message.addEventListener('touchmove', function() { clearTimeout(longPressTimer); });
                
                // Click event for desktop
                 message.addEventListener('click', function(e) {
                     const textToCopy = this.innerText.replace('Copied!', '').trim();
                     copyToClipboard(textToCopy);
                     const tooltip = this.querySelector('.copy-tooltip');
                     if (tooltip) {
                         tooltip.style.display = 'block';
                         setTimeout(() => { tooltip.style.display = 'none'; }, 1500);
                     }
                 });
            });
            
            // Collapsible About Us section
            const aboutUsHeader = document.querySelector('.about-us-header');
            const aboutUsContent = document.querySelector('.about-us-content');
            if (aboutUsHeader && aboutUsContent) {
                // Initialize state if not already done
                if (!aboutUsContent.classList.contains('initialized')) {
                     aboutUsContent.style.display = 'none'; // Start collapsed
                     aboutUsContent.classList.add('initialized');
                }
                // Toggle visibility on header click
                aboutUsHeader.addEventListener('click', function() {
                    aboutUsContent.style.display = (aboutUsContent.style.display === 'none') ? 'block' : 'none';
                });
            }
        }, 1000); // Delay helps ensure elements exist
    });
    </script>
    """

def apply_custom_css():
    """Injects custom CSS into the Streamlit app."""
    css = get_custom_css()
    st.markdown(css, unsafe_allow_html=True)

def add_javascript_functionality():
    """Injects custom JavaScript into the Streamlit app."""
    js = get_custom_javascript()
    st.markdown(js, unsafe_allow_html=True)

# --- Page Configuration --- 
st.set_page_config(page_title="DeepHydro AI Forecasting", layout="wide")
# Apply CSS and JS early in the script execution
apply_custom_css()
add_javascript_functionality() 

# --- Capture User Agent --- 
def capture_user_agent():
    """Capture and store the user agent in session state using components.html."""
    if 'user_agent' not in st.session_state:
        try:
            # The 'key' argument here is for the Streamlit component itself, not the postMessage.
            # The postMessage uses its own 'key' ('user_agent_capture_component') to update session_state.
            components.html(
                """
                <script>
                // Send the user agent back to Streamlit via postMessage
                window.parent.postMessage({
                    isStreamlitMessage: true,
                    type: "streamlit:setComponentValue",
                    key: "user_agent_capture_component", /* Key used by Streamlit to retrieve value */
                    value: navigator.userAgent
                }, "*");
                </script>
                """,
                height=0,
                key="user_agent_capture_component_html" # Unique key for the component
            )
            # Check if the value was set in session state by the postMessage callback
            if 'user_agent_capture_component' in st.session_state and st.session_state.user_agent_capture_component:
                 st.session_state.user_agent = st.session_state.user_agent_capture_component
            else:
                 st.session_state.user_agent = "Unknown (Capture Pending)"
        except Exception as e:
            st.session_state.user_agent = "Unknown (Capture Failed)"

# --- Initialize Firebase and User Profile --- 
firebase_initialized = initialize_firebase()
capture_user_agent() # Attempt to capture user agent

# Initialize user profile
if 'user_profile' not in st.session_state:
    if firebase_initialized:
        user_id = get_persistent_user_id()
        st.session_state.user_profile, _ = get_or_create_user_profile(user_id)
    else:
        st.session_state.user_profile = None

# --- Gemini API Configuration --- 
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")
gemini_configured = False
if GEMINI_API_KEY and GEMINI_API_KEY != "Gemini_api_key":
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        generation_config = genai.types.GenerationConfig(temperature=0.7, top_p=0.95, top_k=40, max_output_tokens=4000)
        gemini_model_report = genai.GenerativeModel(model_name="gemini-2.5-flash-preview-04-17", generation_config=generation_config)
        gemini_model_chat = genai.GenerativeModel(model_name="gemini-2.5-flash-preview-04-17", generation_config=generation_config)
        gemini_configured = True
    except Exception as e:
        st.error(f"Error configuring Gemini API: {e}. AI features might be limited.")
else:
    st.warning("Gemini API Key not found or is placeholder. AI features will be disabled. Set GOOGLE_API_KEY environment variable.")

# --- Model Paths & Constants --- 
STANDARD_MODEL_PATH = "standard_model.h5"
STANDARD_MODEL_SEQUENCE_LENGTH = 60
if os.path.exists(STANDARD_MODEL_PATH):
    try:
        _std_model_temp = tf.keras.models.load_model(STANDARD_MODEL_PATH, compile=False)
        STANDARD_MODEL_SEQUENCE_LENGTH = _std_model_temp.input_shape[1]
        del _std_model_temp
    except Exception as e:
        st.warning(f"Could not load standard model from {STANDARD_MODEL_PATH}: {e}. Using default {STANDARD_MODEL_SEQUENCE_LENGTH}.")
else:
    st.warning(f"Standard model file not found: {STANDARD_MODEL_PATH}.")

# --- Helper Functions (Data, Model, Prediction) --- 
@st.cache_data
def load_and_clean_data(uploaded_file_content):
    try:
        df = pd.read_excel(io.BytesIO(uploaded_file_content), engine="openpyxl")
        if df.shape[1] < 2: st.error("File must have at least Date and Level columns."); return None
        
        # Prioritize exact column names "Date" and "Level" as requested
        date_col_found = "Date" in df.columns
        level_col_found = "Level" in df.columns

        if not date_col_found and not level_col_found:
            # Fallback to intelligent guessing if explicit names not found
            date_col = next((col for col in df.columns if any(kw in str(col).lower() for kw in ["date", "time"]) and pd.api.types.is_datetime64_any_dtype(df[col])), 
                            next((col for col in df.columns if any(kw in str(col).lower() for kw in ["date", "time"]) ), None))
            level_col = next((col for col in df.columns if any(kw in str(col).lower() for kw in ["level", "groundwater", "gwl", "value"]) and pd.api.types.is_numeric_dtype(df[col])), 
                            next((col for col in df.columns if any(kw in str(col).lower() for kw in ["level", "groundwater", "gwl", "value"]) ), None))
            
            if not date_col: st.error("Cannot find a suitable Date/Time column. Please ensure the first column is named 'Date'."); return None
            if not level_col: st.error("Cannot find a suitable Level/Value column. Please ensure the second column is named 'Level'."); return None
            
            st.success(f"Identified columns (auto-detected): Date='{date_col}', Level='{level_col}'.")
            df = df.rename(columns={date_col: "Date", level_col: "Level"})[["Date", "Level"]]
        elif not date_col_found:
            st.error("The first column must be named 'Date'.")
            return None
        elif not level_col_found:
            st.error("The second column must be named 'Level'.")
            return None
        else:
            df = df[["Date", "Level"]] # Explicitly use "Date" and "Level" columns
            st.success("Identified columns: Date='Date', Level='Level'.")

        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        df["Level"] = pd.to_numeric(df["Level"], errors="coerce")
        
        initial_rows = len(df)
        df.dropna(subset=["Date", "Level"], inplace=True)
        if len(df) < initial_rows: st.warning(f"Dropped {initial_rows - len(df)} rows with invalid data.")
        if df.empty: st.error("No valid data remaining."); return None
        
        df = df.sort_values(by="Date").reset_index(drop=True)
        if df.duplicated(subset=["Date"]).any():
            duplicates_count = df.duplicated(subset=["Date"]).sum()
            st.warning(f"Found {duplicates_count} duplicate dates. Keeping first occurrence.")
            df = df.drop_duplicates(subset=["Date"], keep="first")
            
        if df["Level"].isnull().any():
            missing_before = df["Level"].isnull().sum()
            df["Level"] = df["Level"].interpolate(method="linear", limit_direction="both")
            st.warning(f"Filled {missing_before} missing level values using interpolation.")
        if df["Level"].isnull().any(): st.error("Could not fill all missing values."); return None
        
        st.success("Data loaded and cleaned!")
        return df
    except Exception as e: st.error(f"Error loading/cleaning data: {e}"); return None

def create_sequences(data, sequence_length):
    X, y = [], []
    for i in range(len(data) - sequence_length):
        X.append(data[i:(i + sequence_length)])
        y.append(data[i + sequence_length])
    return np.array(X), np.array(y)

@st.cache_resource
def load_keras_model_from_file(uploaded_file_obj, model_name_for_log):
    temp_model_path = f"temp_{model_name_for_log.replace(' ', '_')}.h5"
    try:
        with open(temp_model_path, "wb") as f: f.write(uploaded_file_obj.getbuffer())
        model = tf.keras.models.load_model(temp_model_path, compile=False)
        sequence_length = model.input_shape[1]
        st.success(f"Loaded {model_name_for_log}. Sequence length: {sequence_length}")
        return model, sequence_length
    except Exception as e: st.error(f"Error loading Keras model {model_name_for_log}: {e}"); return None, None
    finally: 
        if os.path.exists(temp_model_path): os.remove(temp_model_path)

@st.cache_resource
def load_standard_model_cached(path):
    try:
        model = tf.keras.models.load_model(path, compile=False)
        sequence_length = model.input_shape[1]
        return model, sequence_length
    except Exception as e: st.error(f"Error loading standard Keras model from {path}: {e}"); return None, None

def build_lstm_model(sequence_length, n_features=1):
    model = Sequential([
        LSTM(40, activation="relu", input_shape=(sequence_length, n_features)), 
        Dropout(0.5), 
        Dense(1)
    ])
    model.compile(optimizer="adam", loss="mean_squared_error")
    return model

def predict_with_dropout_uncertainty(model, last_sequence_scaled, n_steps, n_iterations, scaler, model_sequence_length):
    all_predictions = []
    current_sequence = last_sequence_scaled.copy().reshape(1, model_sequence_length, 1)
    
    @tf.function
    def predict_step_training_true(inp):
        return model(inp, training=True)
        
    progress_bar = st.progress(0)
    status_text = st.empty()
    status_text.text(f"Running MC Dropout (0/{n_iterations})...")
    
    for i in range(n_iterations):
        iteration_predictions_scaled = []
        temp_sequence = current_sequence.copy()
        for _ in range(n_steps):
            next_pred_scaled = predict_step_training_true(temp_sequence).numpy()[0,0]
            iteration_predictions_scaled.append(next_pred_scaled)
            new_step = np.array([[next_pred_scaled]]).reshape(1, 1, 1)
            temp_sequence = np.append(temp_sequence[:, 1:, :], new_step, axis=1)
            
        all_predictions.append(iteration_predictions_scaled)
        progress_percentage = (i + 1) / n_iterations
        progress_bar.progress(progress_percentage)
        status_text.text(f"Running MC Dropout ({i+1}/{n_iterations})...")
        
    progress_bar.empty(); status_text.empty()
    
    predictions_array_scaled = np.array(all_predictions)
    mean_preds_scaled = np.mean(predictions_array_scaled, axis=0)
    std_devs_scaled = np.std(predictions_array_scaled, axis=0)
    
    mean_preds = scaler.inverse_transform(mean_preds_scaled.reshape(-1, 1)).flatten()
    
    ci_multiplier = 1.96
    lower_bound_scaled = mean_preds_scaled - ci_multiplier * std_devs_scaled
    upper_bound_scaled = mean_preds_scaled + ci_multiplier * std_devs_scaled
    lower_bound = scaler.inverse_transform(lower_bound_scaled.reshape(-1, 1)).flatten()
    upper_bound = scaler.inverse_transform(upper_bound_scaled.reshape(-1, 1)).flatten()
    
    min_uncertainty_percent = 0.05 
    for i in range(len(mean_preds)):
        if mean_preds[i] != 0:
            min_uncertainty_value = abs(mean_preds[i] * min_uncertainty_percent / 2.0)
            current_half_range = (upper_bound[i] - lower_bound[i]) / 2.0
            if current_half_range < min_uncertainty_value:
                lower_bound[i] = mean_preds[i] - min_uncertainty_value
                upper_bound[i] = mean_preds[i] + min_uncertainty_value
        else:
             abs_uncertainty = 0.01
             lower_bound[i] = -abs_uncertainty
             upper_bound[i] = abs_uncertainty
             
    return mean_preds, lower_bound, upper_bound

def calculate_metrics(y_true, y_pred):
    if not isinstance(y_true, np.ndarray): y_true = np.array(y_true)
    if not isinstance(y_pred, np.ndarray): y_pred = np.array(y_pred)
    if len(y_true) == 0 or len(y_pred) == 0 or len(y_true) != len(y_pred):
        return {"RMSE": np.nan, "MAE": np.nan, "MAPE": np.nan}
    
    mask = np.isfinite(y_true) & np.isfinite(y_pred)
    if not np.any(mask): return {"RMSE": np.nan, "MAE": np.nan, "MAPE": np.nan}
    y_true_clean, y_pred_clean = y_true[mask], y_pred[mask]
    if len(y_true_clean) == 0: return {"RMSE": np.nan, "MAE": np.nan, "MAPE": np.nan}
    
    rmse = np.sqrt(mean_squared_error(y_true_clean, y_pred_clean))
    mae = mean_absolute_error(y_true_clean, y_pred_clean)
    
    mape_mask = (y_true_clean != 0)
    mape = mean_absolute_percentage_error(y_true_clean[mape_mask], y_pred_clean[mape_mask]) * 100 if np.any(mape_mask) else np.nan
        
    return {"RMSE": rmse, "MAE": mae, "MAPE": mape}

# --- Plotting Functions --- 
def create_forecast_plot(historical_df, forecast_df):
    fig = go.Figure()
    # Use Streamlit's primary color for historical data for theme compatibility
    fig.add_trace(go.Scatter(x=historical_df["Date"], y=historical_df["Level"], mode="lines", name="Historical Data", line=dict(color="rgb(30, 144, 255)"))) # DodgerBlue as a primary example
    fig.add_trace(go.Scatter(x=forecast_df["Date"], y=forecast_df["Forecast"], mode="lines", name="Forecast", line=dict(color="rgb(255, 127, 14)"))) # Orange for forecast
    fig.add_trace(go.Scatter(x=forecast_df["Date"], y=forecast_df["Upper_CI"], mode="lines", name="Upper CI (95%)", line=dict(width=0), showlegend=False))
    fig.add_trace(go.Scatter(x=forecast_df["Date"], y=forecast_df["Lower_CI"], mode="lines", name="Lower CI (95%)", line=dict(width=0), fillcolor="rgba(255, 127, 14, 0.2)", fill="tonexty", showlegend=True))
    fig.update_layout(title="Groundwater Level: Historical Data & AI Forecast", xaxis_title="Date", yaxis_title="Groundwater Level", hovermode="x unified", legend=dict(yanchor="top", y=0.99, xanchor="left", x=0.01), template="plotly_white")
    return fig

def create_loss_plot(history_dict):
    if not history_dict or not isinstance(history_dict, dict) or "loss" not in history_dict or "val_loss" not in history_dict:
        fig = go.Figure().update_layout(title="No Training History Available", xaxis_title="Epoch", yaxis_title="Loss")
        fig.add_annotation(text="Training history unavailable.",xref="paper", yref="paper", x=0.5, y=0.5, showarrow=False)
        return fig
    history_df = pd.DataFrame(history_dict); history_df["Epoch"] = history_df.index + 1
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=history_df["Epoch"], y=history_df["loss"], mode="lines", name="Training Loss", line=dict(color="rgb(30, 144, 255)"))) # DodgerBlue
    fig.add_trace(go.Scatter(x=history_df["Epoch"], y=history_df["val_loss"], mode="lines", name="Validation Loss", line=dict(color="rgb(255, 127, 14)"))) # Orange
    fig.update_layout(title="Model Training & Validation Loss", xaxis_title="Epoch", yaxis_title="Loss (MSE)", hovermode="x unified", template="plotly_white")
    return fig

# --- Gemini API Functions --- 
def generate_gemini_report(hist_df, forecast_df, metrics, language):
    if not gemini_configured: return "AI report disabled. Configure Gemini API Key."
    if hist_df is None or forecast_df is None or metrics is None: return "Error: Insufficient data for AI report."
    try:
        prompt = f"""Act as a professional hydrologist. Provide a concise report in {language} based on the historical data, forecast, and metrics. Focus on trends, forecast reliability (mentioning C.I.), implications, and recommendations. **IMPORTANT: Do NOT discuss technical model details (architecture, training). Focus on data and outcomes.**

Historical Summary:
{hist_df["Level"].describe().to_string()}

Forecast Summary:
{forecast_df[["Forecast", "Lower_CI", "Upper_CI"]].describe().to_string()}

Metrics:
RMSE: {metrics.get('RMSE', 'N/A'):.4f}
MAE: {metrics.get('MAE', 'N/A'):.4f}
MAPE: {metrics.get('MAPE', 'N/A'):.2f}%

Generate the report:"""
        response = gemini_model_report.generate_content(prompt)
        forbidden_terms = ["lstm", "long short-term memory", "epoch", "layer", "dropout", "adam optimizer", "sequence length"]
        cleaned_text = response.text
        for term in forbidden_terms: cleaned_text = cleaned_text.replace(term, "[modeling technique]")
        return cleaned_text
    except Exception as e: st.error(f"Error generating AI report: {e}"); return f"Error: {e}"

def get_gemini_chat_response(user_query, chat_hist, hist_df, forecast_df, metrics, ai_report):
    if not gemini_configured: return "AI chat disabled. Configure Gemini API Key."
    if hist_df is None or forecast_df is None or metrics is None: return "Error: Insufficient context for AI chat."
    try:
        # Refined persona and removal of intro phrases
        context_parts = [
            "You are an experienced hydrogeologist and data analyst with over 10 years of experience in groundwater analysis and predictive modeling.",
            "Your task is to interpret the following groundwater data and forecasts.",
            "**IMPORTANT: Do NOT discuss internal AI model mechanics. Focus only on data interpretation, patterns, trends, uncertainties, and implications for groundwater behavior.**",
            "**Do NOT add any introductory phrases before your answer.** Start directly with the analysis.",
            "",
            "### Historical Groundwater Summary:",
            hist_df["Level"].describe().to_string(),
            "",
            "### Forecast Results Summary:",
            forecast_df[["Forecast", "Lower_CI", "Upper_CI"]].describe().to_string(),
            "",
            f"### Forecast Accuracy Metrics:\nRMSE = {metrics.get('RMSE', 'N/A'):.4f}\nMAE = {metrics.get('MAE', 'N/A'):.4f}\nMAPE = {metrics.get('MAPE', 'N/A'):.2f}%",
            "",
            "### Existing AI Report (if available):",
            ai_report if ai_report else "(No prior AI report available.)",
            "",
            "### Instructions:",
            "- Provide insights from an experienced hydrogeologist and data analyst (10 years experience).",
            "- Identify trends, anomalies, and seasonal patterns in groundwater level data.",
            "- Explain the forecast ranges (CI) and what they imply for groundwater conditions.",
            "- Compare historical and forecast data to infer changes in groundwater behavior.",
            "- Suggest potential implications for resource management, policy, or risk.",
            "- Structure your response directly, without an introduction.",
            "- Keep responses conversational but highly professional and data-driven, like an expert providing insights.",
            "",
            "### Previous Conversation (for context):"
        ]
        for sender, message in chat_hist[-6:]: context_parts.append(f"{sender}: {message}")
        context_parts.append(f"User: {user_query}"); context_parts.append("AI:") 
        context = "\n".join(context_parts)
        
        response = gemini_model_chat.generate_content(context)
        forbidden_terms = ["lstm", "long short-term memory", "epoch", "layer", "dropout", "adam optimizer", "sequence length", "as an AI model", "I am an AI", "hello", "hi", "greetings", "certainly", "absolutely", "of course"]
        cleaned_text = response.text
        for term in forbidden_terms: cleaned_text = cleaned_text.replace(term, "[modeling technique]")
        # Further refine to ensure no intros
        if cleaned_text.lower().startswith(("hello", "hi", "greetings", "as an ai model", "i am an ai", "certainly", "absolutely", "of course")):
            # Find the first sentence and remove it if it sounds like an intro
            sentences = cleaned_text.split('.')
            if len(sentences) > 1:
                first_sentence_lower = sentences[0].strip().lower()
                if any(intro_phrase in first_sentence_lower for intro_phrase in ["hello", "hi", "greetings", "as an ai model", "i am an ai", "certainly", "absolutely", "of course"]):
                    cleaned_text = '.'.join(sentences[1:]).strip()
            else: # If it's a single short sentence intro, try to remove it
                if any(intro_phrase in cleaned_text.lower() for intro_phrase in ["hello", "hi", "greetings", "as an ai model", "i am an ai", "certainly", "absolutely", "of course"]):
                    cleaned_text = "Analysis: " 
        return cleaned_text.strip() 
    except Exception as e: st.error(f"Error in AI chat: {e}"); return f"Error: {e}"

# --- Main Forecasting Pipeline --- 
def run_forecast_pipeline(df, model_choice, forecast_horizon, custom_model_file_obj, 
                        sequence_length_train_param, epochs_train_param, 
                        mc_iterations_param, use_custom_scaler_params_flag, custom_scaler_min_param, custom_scaler_max_param):
    st.info(f"Starting forecast: {model_choice}")
    model, history_data, scaler_obj = None, None, MinMaxScaler(feature_range=(0, 1))
    model_sequence_length = sequence_length_train_param
    
    try:
        st.info("Step 1: Preparing Model...")
        if model_choice == "Standard Pre-trained Model":
            if os.path.exists(STANDARD_MODEL_PATH):
                model, model_sequence_length = load_standard_model_cached(STANDARD_MODEL_PATH)
                if model is None: return None, None, None, None
            else: st.error(f"Standard model not found: {STANDARD_MODEL_PATH}."); return None, None, None, None
        elif model_choice == "Upload Custom .h5 Model" and custom_model_file_obj:
            model, model_sequence_length = load_keras_model_from_file(custom_model_file_obj, "Custom Model")
            if model is None: return None, None, None, None
        elif model_choice == "Train New Model":
            st.info("Training New Model selected. Model will be built and trained.")
        else:
            st.error("Invalid model choice or missing file."); return None, None, None, None
        
        st.session_state.model_sequence_length = model_sequence_length
        st.info(f"Model ready. Sequence length: {model_sequence_length}")

        st.info("Step 2: Preprocessing Data (Scaling)...")
        if use_custom_scaler_params_flag and custom_scaler_min_param is not None and custom_scaler_max_param is not None and custom_scaler_min_param < custom_scaler_max_param:
            scaler_obj.fit(np.array([[custom_scaler_min_param], [custom_scaler_max_param]]))
            scaled_data = scaler_obj.transform(df["Level"].values.reshape(-1, 1))
            st.info(f"Using custom scaler: min={custom_scaler_min_param}, max={custom_scaler_max_param}")
        else:
            scaled_data = scaler_obj.fit_transform(df["Level"].values.reshape(-1, 1))
            st.info(f"Using fitted scaler: min={scaler_obj.data_min_[0]:.4f}, max={scaler_obj.data_max_[0]:.4f}")
        st.info("Scaling complete.")

        st.info(f"Step 3: Creating sequences (length {model_sequence_length})...")
        if len(df) <= model_sequence_length: st.error(f"Not enough data ({len(df)} rows) for sequence length ({model_sequence_length})."); return None, None, None, None
        X, y = create_sequences(scaled_data, model_sequence_length)
        if len(X) == 0: st.error("Could not create sequences."); return None, None, None, None
        st.info(f"Sequences created: {len(X)}")

        evaluation_metrics = {"RMSE": np.nan, "MAE": np.nan, "MAPE": np.nan}
        if model_choice == "Train New Model":
            st.info(f"Step 4a: Training New Model (Epochs: {epochs_train_param})...")
            X_train, X_val, y_train, y_val = train_test_split(X, y, test_size=0.2, shuffle=False)
            if len(X_train) == 0 or len(X_val) == 0: st.error("Not enough data for train/validation split."); return None, None, None, None
            
            model = build_lstm_model(model_sequence_length)
            early_stopping = EarlyStopping(monitor="val_loss", patience=10, restore_best_weights=True)
            train_progress = st.progress(0); status_text_train = st.empty()
            class ProgressCallback(tf.keras.callbacks.Callback):
                def on_epoch_end(self, epoch, logs=None):
                    progress = (epoch + 1) / epochs_train_param
                    train_progress.progress(progress)
                    status_text_train.text(f"Training Epoch {epoch+1}/{epochs_train_param} - Loss: {logs['loss']:.4f}, Val Loss: {logs['val_loss']:.4f}")
            
            history_obj = model.fit(X_train, y_train, epochs=epochs_train_param, batch_size=32, validation_data=(X_val, y_val), callbacks=[early_stopping, ProgressCallback()], verbose=0)
            history_data = history_obj.history
            train_progress.empty(); status_text_train.empty()
            st.success("Training complete.")
            
            st.info("Evaluating trained model...")
            val_predictions_scaled = model.predict(X_val)
            val_predictions = scaler_obj.inverse_transform(val_predictions_scaled)
            y_val_actual = scaler_obj.inverse_transform(y_val)
            evaluation_metrics = calculate_metrics(y_val_actual, val_predictions)
            st.success("Evaluation complete.")
        else: # Pre-trained model evaluation
            st.info("Step 4b: Evaluating Pre-trained Model (Pseudo-Validation)...")
            if len(X) > 5 and model: # Ensure model exists for pre-trained paths
                val_split_idx = max(1, int(len(X) * 0.8))
                X_val_pseudo, y_val_pseudo = X[val_split_idx:], y[val_split_idx:]
                if len(X_val_pseudo) > 0:
                    val_predictions_scaled = model.predict(X_val_pseudo)
                    val_predictions = scaler_obj.inverse_transform(val_predictions_scaled)
                    y_val_actual = scaler_obj.inverse_transform(y_val_pseudo)
                    evaluation_metrics = calculate_metrics(y_val_actual, val_predictions)
                    st.success("Pseudo-evaluation complete.")
                else: st.warning("Not enough data for pseudo-validation.")
            else: st.warning("Not enough sequences for pseudo-validation or model not loaded.")


        st.info(f"Step 5: Forecasting {forecast_horizon} Steps (MC Dropout: {mc_iterations_param})...")
        if model is None: # Ensure model is available for prediction
            st.error("Model not available for forecasting. Please check model loading/training steps.")
            return None, None, None, None

        last_sequence_scaled_for_pred = scaled_data[-model_sequence_length:]
        mean_forecast, lower_bound, upper_bound = predict_with_dropout_uncertainty(
            model, last_sequence_scaled_for_pred, forecast_horizon, mc_iterations_param, scaler_obj, model_sequence_length
        )
        st.success("Forecasting complete.")

        last_date = df["Date"].iloc[-1]
        try: freq = pd.infer_freq(df["Date"].dropna()) or "D"
        except Exception: freq = "D"
        try: date_offset = pd.tseries.frequencies.to_offset(freq)
        except ValueError: 
            st.warning(f"Invalid frequency '{freq}'. Defaulting to daily ('D').")
            date_offset = pd.DateOffset(days=1); freq = 'D'
            
        forecast_dates = pd.date_range(start=last_date + date_offset, periods=forecast_horizon, freq=freq)
        forecast_df = pd.DataFrame({"Date": forecast_dates, "Forecast": mean_forecast, "Lower_CI": lower_bound, "Upper_CI": upper_bound})
        
        st.info("Forecast pipeline finished.")
        return forecast_df, evaluation_metrics, history_data, scaler_obj

    except Exception as e:
        st.error(f"Error in forecast pipeline: {e}")
        import traceback; st.error(traceback.format_exc())
        return None, None, None, None

# --- New: Data Analysis & Statistics Tab Content ---
def render_data_analysis_tab(cleaned_data, forecast_results):
    st.header("Data Analysis & Statistics")

    if cleaned_data is None:
        st.info("Please upload data using the sidebar to see analysis.")
        return
    
    st.subheader("Historical Data Overview")
    st.dataframe(cleaned_data.describe().transpose())
    st.write("This table provides a statistical summary of the uploaded historical groundwater level data, including count, mean, standard deviation, min/max values, and quartiles.")

    fig_hist = px.histogram(cleaned_data, x="Level", nbins=30, 
                            title="Distribution of Historical Groundwater Levels",
                            labels={"Level": "Groundwater Level"},
                            template="plotly_white")
    st.plotly_chart(fig_hist, use_container_width=True)
    st.write("This histogram shows the frequency distribution of historical groundwater levels, helping to visualize common ranges and potential skewness.")

    # Autocorrelation Plot for Historical Data
    st.subheader("Autocorrelation of Historical Data")
    try:
        # Use Matplotlib for plot_acf, then convert to Plotly for consistency
        fig_acf_mpl, ax_acf = plt.subplots(figsize=(10, 5))
        plot_acf(cleaned_data["Level"], ax=ax_acf, lags=min(len(cleaned_data)//2 - 1, 50), title="Autocorrelation Function (ACF) of Historical Groundwater Levels")
        ax_acf.set_xlabel("Lags")
        ax_acf.set_ylabel("Autocorrelation Coefficient")
        plt.tight_layout()
        st.pyplot(fig_acf_mpl) # Streamlit can directly render Matplotlib figures
        plt.close(fig_acf_mpl) # Close the Matplotlib figure to free memory
        st.write("The Autocorrelation Function (ACF) plot helps identify patterns in the data over time, such as seasonality (peaks at regular intervals) or trends (slow decay).")
    except Exception as e:
        st.warning(f"Could not generate Autocorrelation plot: {e}. Ensure enough data points are available.")
    
    # Rolling Statistics for Historical Data
    st.subheader("Rolling Statistics of Historical Data")
    if len(cleaned_data) > 30: # Need enough data for rolling window
        fig_rolling = go.Figure()
        window_size = min(30, len(cleaned_data) - 1)
        rolling_mean = cleaned_data["Level"].rolling(window=window_size).mean()
        rolling_std = cleaned_data["Level"].rolling(window=window_size).std()

        fig_rolling.add_trace(go.Scatter(x=cleaned_data["Date"], y=cleaned_data["Level"], mode="lines", name="Original Level", line=dict(color="lightgray", width=0.5), opacity=0.7))
        fig_rolling.add_trace(go.Scatter(x=cleaned_data["Date"], y=rolling_mean, mode="lines", name=f"Rolling Mean ({window_size} days)", line=dict(color="rgb(30, 144, 255)", width=2)))
        fig_rolling.add_trace(go.Scatter(x=cleaned_data["Date"], y=rolling_std, mode="lines", name=f"Rolling Std Dev ({window_size} days)", line=dict(color="rgb(255, 127, 14)", dash='dash', width=1.5)))
        fig_rolling.update_layout(title=f"Rolling Mean and Standard Deviation (Window: {window_size} days)",
                                  xaxis_title="Date", yaxis_title="Groundwater Level",
                                  hovermode="x unified", template="plotly_white")
        st.plotly_chart(fig_rolling, use_container_width=True)
        st.write("This plot shows the moving average and standard deviation, highlighting trends and changes in data variability over time.")
    else:
        st.info("Not enough data for rolling statistics plot (requires at least 30 data points).")

    # Monthly Box Plot (if data spans multiple years)
    if cleaned_data['Date'].dt.year.nunique() > 1:
        st.subheader("Monthly Groundwater Level Distribution")
        cleaned_data['Month'] = cleaned_data['Date'].dt.month_name()
        month_order = ["January", "February", "March", "April", "May", "June", 
                       "July", "August", "September", "October", "November", "December"]
        cleaned_data['Month'] = pd.Categorical(cleaned_data['Month'], categories=month_order, ordered=True)
        
        fig_monthly_box = px.box(cleaned_data, x="Month", y="Level", 
                                 title="Distribution of Groundwater Levels by Month",
                                 labels={"Level": "Groundwater Level"},
                                 template="plotly_white")
        fig_monthly_box.update_layout(xaxis_title="Month", yaxis_title="Groundwater Level")
        st.plotly_chart(fig_monthly_box, use_container_width=True)
        st.write("This box plot shows the distribution of groundwater levels for each month across all years, indicating seasonal patterns and variability.")

    if forecast_results is not None:
        st.subheader("Forecasted Data Overview")
        st.dataframe(forecast_results[['Forecast', 'Lower_CI', 'Upper_CI']].describe().transpose())
        st.write("This table provides a statistical summary of the forecasted groundwater levels, including the mean forecast and confidence intervals.")

        fig_forecast_hist = px.histogram(forecast_results, x="Forecast", nbins=30,
                                         title="Distribution of Forecasted Groundwater Levels",
                                         labels={"Forecast": "Forecasted Level"},
                                         template="plotly_white", color_discrete_sequence=[px.colors.qualitative.Plotly[1]]) # Orange color
        st.plotly_chart(fig_forecast_hist, use_container_width=True)
        st.write("This histogram illustrates the distribution of the predicted groundwater levels over the forecast horizon.")

        st.subheader("Historical vs. Forecasted Data Comparison")
        # Combine data for comparison plot
        combined_df = pd.DataFrame({
            'Level': pd.concat([cleaned_data['Level'], forecast_results['Forecast']]),
            'Type': ['Historical'] * len(cleaned_data) + ['Forecast'] * len(forecast_results)
        })
        fig_combined_dist = px.histogram(combined_df, x="Level", color="Type", histnorm='probability density',
                                        barmode='overlay', opacity=0.7,
                                        title="Density Distribution: Historical vs. Forecasted Levels",
                                        labels={"Level": "Groundwater Level", "count": "Density"},
                                        color_discrete_map={'Historical': "rgb(30, 144, 255)", 'Forecast': "rgb(255, 127, 14)"},
                                        template="plotly_white")
        fig_combined_dist.update_layout(bargap=0.1)
        st.plotly_chart(fig_combined_dist, use_container_width=True)
        st.write("This overlaid histogram compares the density distributions of historical and forecasted groundwater levels, showing how the predicted levels align or differ from past observations.")
        
        st.subheader("Key Statistical Differences")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Historical Mean", f"{cleaned_data['Level'].mean():.2f}")
            st.metric("Forecast Mean", f"{forecast_results['Forecast'].mean():.2f}")
        with col2:
            st.metric("Historical Std Dev", f"{cleaned_data['Level'].std():.2f}")
            st.metric("Forecast Std Dev", f"{forecast_results['Forecast'].std():.2f}")
        with col3:
            st.metric("Historical Range", f"{cleaned_data['Level'].max() - cleaned_data['Level'].min():.2f}")
            st.metric("Forecast Range", f"{forecast_results['Forecast'].max() - forecast_results['Forecast'].min():.2f}")
        st.write("These metrics highlight the central tendency, spread, and overall range differences between the historical and forecasted groundwater levels.")

    else:
        st.info("Run a forecast to see combined data analysis.")

# --- Initialize Session State --- 
def initialize_session_state():
    defaults = {
        "cleaned_data": None, "forecast_results": None, "evaluation_metrics": None, 
        "training_history": None, "ai_report": None, "scaler_object": None, 
        "forecast_plot_fig": None, "uploaded_data_filename": None,
        "active_tab": 0, "report_language": "English", "chat_history": [], 
        "chat_active": False, "model_sequence_length": STANDARD_MODEL_SEQUENCE_LENGTH, 
        "run_forecast_triggered": False, "about_us_expanded": False,
        "persistent_user_id": None, # Will be set by get_persistent_user_id
        "user_profile": None, # Will be set by get_or_create_user_profile
        "admin_authenticated": False, 
        "session_visit_logged": False, "user_agent": None
    }
    for key, default_value in defaults.items():
        if key not in st.session_state: st.session_state[key] = default_value

initialize_session_state()

# --- Sidebar --- 
with st.sidebar:
    st.title("DeepHydro AI")
    
    # Removed User Info / Logout section as Google Sign-in is removed.
    
    if firebase_initialized: log_visitor_activity("Sidebar", "view")
    
    st.header("1. Upload Data")
    uploaded_data_file = st.file_uploader("Choose XLSX data file", type="xlsx", key="data_uploader")
    st.info("Note: Your XLSX file should have the first column named 'Date' and the second column named 'Level'.")
    
    st.header("2. Model & Forecast")
    model_choice = st.selectbox("Model Type", ("Standard Pre-trained Model", "Train New Model", "Upload Custom .h5 Model"), key="model_select")
    
    if firebase_initialized:
        if 'last_model_choice' not in st.session_state or st.session_state.last_model_choice != model_choice:
             log_visitor_activity("Sidebar", "select_model", feature_used=model_choice)
             st.session_state.last_model_choice = model_choice

    custom_model_file_obj_sidebar = None
    custom_scaler_min_sidebar, custom_scaler_max_sidebar = None, None
    use_custom_scaler_sidebar = False
    default_sequence_length = st.session_state.get("model_sequence_length", STANDARD_MODEL_SEQUENCE_LENGTH)
    sequence_length_train_sidebar = default_sequence_length
    epochs_train_sidebar = 50

    if model_choice == "Upload Custom .h5 Model":
        custom_model_file_obj_sidebar = st.file_uploader("Upload .h5 model", type="h5", key="custom_h5_uploader")
        use_custom_scaler_sidebar = st.checkbox("Provide custom scaler params?", value=False, key="use_custom_scaler_cb")
        if use_custom_scaler_sidebar:
            st.markdown("Enter **original min/max** values:")
            custom_scaler_min_sidebar = st.number_input("Original Min", value=0.0, format="%.4f", key="custom_scaler_min_in")
            custom_scaler_max_sidebar = st.number_input("Original Max", value=1.0, format="%.4f", key="custom_scaler_max_in")
    elif model_choice == "Standard Pre-trained Model":
        st.info(f"Using standard model (Seq Len: {default_sequence_length})")
        use_custom_scaler_sidebar = st.checkbox("Provide custom scaler params?", value=False, key="use_std_scaler_cb")
        if use_custom_scaler_sidebar:
            st.markdown("Enter **original min/max** values:")
            custom_scaler_min_sidebar = st.number_input("Original Min", value=0.0, format="%.4f", key="std_scaler_min_in")
            custom_scaler_max_sidebar = st.number_input("Original Max", value=1.0, format="%.4f", key="std_scaler_max_in")
    elif model_choice == "Train New Model":
        try:
            sequence_length_train_sidebar = st.number_input("Model Sequence Length", min_value=10, max_value=365, value=default_sequence_length, step=10, key="seq_len_train_in")
        except Exception as e:
            st.warning(f"Using default sequence length {default_sequence_length}: {e}")
            sequence_length_train_sidebar = default_sequence_length
        epochs_train_sidebar = st.number_input("Training Epochs", min_value=10, max_value=500, value=50, step=10, key="epochs_train_in")

    mc_iterations_sidebar = st.number_input("MC Dropout Iterations (C.I.)", min_value=20, max_value=500, value=100, step=10, key="mc_iter_in")
    forecast_horizon_sidebar = st.number_input("Forecast Horizon (steps)", min_value=1, max_value=100, value=12, step=1, key="horizon_in")

    # Run Forecast Button (clicking this will switch to Forecast Results tab)
    if st.button("Run Forecast", key="run_forecast_main_btn", use_container_width=True):
        st.session_state.run_forecast_triggered = True
        if st.session_state.cleaned_data is not None:
            if model_choice == "Upload Custom .h5 Model" and not custom_model_file_obj_sidebar:
                st.error("Please upload a custom .h5 model file.")
                st.session_state.run_forecast_triggered = False
            else:
                if firebase_initialized: log_visitor_activity("Sidebar", "run_forecast", feature_used='Forecast')
                with st.spinner(f"Running forecast ({model_choice})..."):
                    forecast_df, metrics, history, scaler_obj = run_forecast_pipeline(
                        st.session_state.cleaned_data, model_choice, forecast_horizon_sidebar, 
                        custom_model_file_obj_sidebar, sequence_length_train_sidebar, epochs_train_sidebar, 
                        mc_iterations_sidebar, use_custom_scaler_sidebar, custom_scaler_min_sidebar, custom_scaler_max_sidebar
                    )
                st.session_state.forecast_results = forecast_df
                st.session_state.evaluation_metrics = metrics
                st.session_state.training_history = history
                st.session_state.scaler_object = scaler_obj
                
                if forecast_df is not None and metrics is not None:
                    st.session_state.forecast_plot_fig = create_forecast_plot(st.session_state.cleaned_data, forecast_df)
                    st.success("Forecast complete! Results updated.")
                    st.session_state.ai_report = None; st.session_state.chat_history = []; st.session_state.chat_active = False
                    st.session_state.active_tab = 1 # Switch to Forecast Results tab
                    st.rerun()
                else:
                    st.error("Forecast pipeline failed. Check messages.")
                    st.session_state.forecast_results = None; st.session_state.evaluation_metrics = None
                    st.session_state.training_history = None; st.session_state.forecast_plot_fig = None
        else:
            st.error("Please upload data first.")
            st.session_state.run_forecast_triggered = False


    st.header("3. AI Analysis")
    st.session_state.report_language = st.selectbox("Report Language", ["English", "French"], key="report_lang_select", disabled=not gemini_configured)
    
    # Generate AI Report Button (clicking this will switch to AI Report tab)
    if st.button("Generate AI Report", key="show_report_btn", disabled=not gemini_configured, use_container_width=True):
        if not gemini_configured: st.error("AI Report disabled. Configure Gemini API Key.")
        elif st.session_state.cleaned_data is not None and st.session_state.forecast_results is not None and st.session_state.evaluation_metrics is not None:
            if firebase_initialized: log_visitor_activity("Sidebar", "generate_report", feature_used='AI Report')
            with st.spinner(f"Generating AI report ({st.session_state.report_language})..."):
                st.session_state.ai_report = generate_gemini_report(
                    st.session_state.cleaned_data, st.session_state.forecast_results,
                    st.session_state.evaluation_metrics, st.session_state.report_language
                )
            if st.session_state.ai_report and not st.session_state.ai_report.startswith("Error:"):
                st.success("AI report generated.")
                st.session_state.active_tab = 4 # Switch to AI Report tab
                st.rerun()
            else: st.error(f"Failed to generate AI report. {st.session_state.ai_report}")
        else: st.error("Data, forecast, and metrics needed. Run forecast first.")

    # AI Assistant Activate Chat Button (clicking this will switch to AI Chatbot tab)
    chat_button_label = "Deactivate Chat" if st.session_state.chat_active else "Activate Chat"
    if st.button(chat_button_label, key="chat_ai_btn_sidebar", use_container_width=True, disabled=not gemini_configured):
        if st.session_state.chat_active:
            st.session_state.chat_active = False; st.session_state.chat_history = []
            if firebase_initialized: log_visitor_activity("Sidebar", "deactivate_chat")
            st.rerun()
        else:
            st.session_state.chat_active = True; st.session_state.active_tab = 5 # Switch to AI Chatbot tab
            if firebase_initialized: log_visitor_activity("Sidebar", "activate_chat", feature_used='AI Chat')
            st.rerun()

    # About Us
    st.markdown('<div class="about-us-header">👥 About Us</div>', unsafe_allow_html=True)
    st.markdown('<div class="about-us-content">', unsafe_allow_html=True)
    st.markdown("Specializing in groundwater forecasting using AI.") 
    st.markdown("**Contact:** [deephydro@example.com](mailto:deephydro@example.com)")
    st.markdown("© 2025 DeepHydro AI Team")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Admin Access (clicking this will switch to Admin Analytics tab)
    st.header("5. Admin")
    if st.button("Analytics Dashboard", key="admin_analytics_btn_sidebar", use_container_width=True):
        if firebase_initialized: log_visitor_activity("Sidebar", "access_admin")
        st.session_state.active_tab = 6 # Switch to Admin Analytics tab
        st.rerun()

# --- Main Application Area --- 
st.markdown("<h1 class='blue-text'>DeepHydro AI Forecasting</h1>", unsafe_allow_html=True)
if firebase_initialized: log_visitor_activity("Main Page", "view")

# App Introduction
st.markdown(f"""
<div class="app-intro">
    <h3>Revolutionizing Water Management with <span class="blue-text">AI</span></h3>
    Unlock insights into future water resources. Our platform leverages advanced <span class="blue-text">Deep learning</span> and <span class="blue-text">Data analysis</span> to provide precise <span class="blue-text">Hydrogeology</span> forecasts, crucial for the <span class="blue-text">protection of national water sources</span> and strategic planning. Experience the power of predictive <span class="blue-text">AI</span> for sustainable water management.
</div>
""", unsafe_allow_html=True)

# Handle data upload
if uploaded_data_file is not None:
    # Corrected the variable name from `uploaded_file` to `uploaded_data_file`
    if st.session_state.get("uploaded_data_filename") != uploaded_data_file.name:
        st.session_state.uploaded_data_filename = uploaded_data_file.name
        with st.spinner("Loading and cleaning data..."): 
            cleaned_df_result = load_and_clean_data(uploaded_data_file.getvalue())
        if cleaned_df_result is not None:
            st.session_state.cleaned_data = cleaned_df_result
            # Reset results
            st.session_state.forecast_results = None; st.session_state.evaluation_metrics = None
            st.session_state.training_history = None; st.session_state.ai_report = None
            st.session_state.chat_history = []; st.session_state.scaler_object = None
            st.session_state.forecast_plot_fig = None
            st.session_state.model_sequence_length = STANDARD_MODEL_SEQUENCE_LENGTH
            st.session_state.run_forecast_triggered = False
            if firebase_initialized: log_visitor_activity("Data Upload", "upload_success")
            st.rerun()
        else:
            st.session_state.cleaned_data = None
            st.error("Data loading failed. Check file format/content.")
            if firebase_initialized: log_visitor_activity("Data Upload", "upload_failure")

# Define tabs (Removed 'key' argument to fix TypeError)
tab_titles = ["Data Preview", "Forecast Results", "Model Evaluation", "Data Analysis & Statistics", "AI Report", "AI Chatbot", "Admin Analytics"]
tabs = st.tabs(tab_titles, active_tab=st.session_state.active_tab) 

# --- Tab Content --- 

# Data Preview Tab
with tabs[0]:
    if firebase_initialized: log_visitor_activity("Tab: Data Preview", "view")
    st.header("Uploaded & Cleaned Data Preview")
    if st.session_state.cleaned_data is not None:
        st.dataframe(st.session_state.cleaned_data)
        st.write(f"Shape: {st.session_state.cleaned_data.shape}")
        col1, col2 = st.columns(2)
        with col1: st.metric("Time Range", f"{st.session_state.cleaned_data['Date'].min():%Y-%m-%d} to {st.session_state.cleaned_data['Date'].max():%Y-%m-%d}")
        with col2: st.metric("Data Points", len(st.session_state.cleaned_data))
        fig_data = go.Figure()
        fig_data.add_trace(go.Scatter(x=st.session_state.cleaned_data["Date"], y=st.session_state.cleaned_data["Level"], mode="lines", name="Level", line=dict(color="rgb(30, 144, 255)"))) # DodgerBlue
        fig_data.update_layout(title="Historical Groundwater Levels", xaxis_title="Date", yaxis_title="Level", template="plotly_white", margin=dict(l=20, r=20, t=40, b=20), height=400)
        st.plotly_chart(fig_data, use_container_width=True)
    else:
        st.info("⬆️ Upload XLSX data using the sidebar.")

# Forecast Results Tab
with tabs[1]:
    if firebase_initialized: log_visitor_activity("Tab: Forecast Results", "view")
    st.header("Forecast Results")
    if st.session_state.forecast_results is not None and isinstance(st.session_state.forecast_results, pd.DataFrame) and not st.session_state.forecast_results.empty:
        if st.session_state.forecast_plot_fig: st.plotly_chart(st.session_state.forecast_plot_fig, use_container_width=True)
        else: st.warning("Forecast plot unavailable.")
        st.subheader("Forecast Data Table")
        st.dataframe(st.session_state.forecast_results, use_container_width=True)

        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            csv_forecast = st.session_state.forecast_results.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="Download Forecast as CSV",
                data=csv_forecast,
                file_name="forecast_results.csv",
                mime="text/csv",
                key="download_forecast_csv",
                use_container_width=True
            )
        with col_dl2:
            excel_buffer = io.BytesIO()
            st.session_state.forecast_results.to_excel(excel_buffer, index=False, engine='openpyxl')
            excel_buffer.seek(0)
            st.download_button(
                label="Download Forecast as XLSX",
                data=excel_buffer,
                file_name="forecast_results.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_forecast_xlsx",
                use_container_width=True
            )

    elif st.session_state.run_forecast_triggered: st.warning("Forecast run attempted, but no results available.")
    else: st.info("Run a forecast (sidebar) to see results.")

# Model Evaluation Tab
with tabs[2]:
    if firebase_initialized: log_visitor_activity("Tab: Model Evaluation", "view")
    st.header("Model Evaluation")
    if st.session_state.evaluation_metrics is not None and isinstance(st.session_state.evaluation_metrics, dict):
        st.subheader("Performance Metrics (Validation/Pseudo-Validation)")
        col1, col2, col3 = st.columns(3)
        rmse_val = st.session_state.evaluation_metrics.get("RMSE", np.nan); mae_val = st.session_state.evaluation_metrics.get("MAE", np.nan); mape_val = st.session_state.evaluation_metrics.get("MAPE", np.nan)
        col1.metric("RMSE", f"{rmse_val:.4f}" if not np.isnan(rmse_val) else "N/A")
        col2.metric("MAE", f"{mae_val:.4f}" if not np.isnan(mae_val) else "N/A")
        col3.metric("MAPE", f"{mape_val:.2f}%" if not np.isnan(mape_val) and mape_val != np.inf else ("N/A" if np.isnan(mape_val) else "Inf"))
        st.subheader("Training Loss (if trained)")
        if st.session_state.training_history:
            loss_fig = create_loss_plot(st.session_state.training_history)
            st.plotly_chart(loss_fig, use_container_width=True)
        else: st.info("No training history (pre-trained model or training failed).")
    elif st.session_state.run_forecast_triggered: st.warning("Forecast run attempted, but no evaluation metrics available.")
    else: st.info("Run a forecast (sidebar) to see evaluation.")

# Data Analysis & Statistics Tab
with tabs[3]:
    if firebase_initialized: log_visitor_activity("Tab: Data Analysis & Statistics", "view")
    render_data_analysis_tab(st.session_state.cleaned_data, st.session_state.forecast_results)

# AI Report Tab
with tabs[4]:
    if firebase_initialized: log_visitor_activity("Tab: AI Report", "view")
    st.header("AI-Generated Scientific Report")
    if not gemini_configured: st.warning("AI features disabled. Configure Gemini API Key.")
    if st.session_state.ai_report: 
        st.markdown(f'<div class="chat-message ai-message">{st.session_state.ai_report}<span class="copy-tooltip">Copied!</span></div>', unsafe_allow_html=True)

        # Download Word Report Button - Moved from Sidebar to Main Area (AI Report Tab)
        st.subheader("Download Report")
        # Check if necessary data for report generation is available
        if st.session_state.forecast_results is not None and \
           st.session_state.evaluation_metrics is not None and \
           st.session_state.ai_report is not None and \
           st.session_state.forecast_plot_fig is not None:
            
            # Create a BytesIO object to store the document in memory
            docx_buffer = io.BytesIO()
            document = Document()
            
            # Set up default font for the document
            style = document.styles['Normal']
            font = style.font
            font.name = 'Calibri' 
            font.size = Pt(11)

            # Title
            document.add_heading('DeepHydro AI Forecasting Report', level=0)
            document.add_paragraph(f"Date Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
            document.add_paragraph("\n") 

            # Forecast Visualization Section
            document.add_heading('1. Forecast Visualization', level=1)
            plot_filename = "forecast_plot.png"
            try:
                if st.session_state.forecast_plot_fig:
                    st.session_state.forecast_plot_fig.write_image(plot_filename, scale=2) 
                    document.add_picture(plot_filename, width=Inches(6.5)) 
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER 
                    document.add_paragraph("\n")
                else:
                    document.add_paragraph("[Forecast plot unavailable]")
            except Exception as img_err:
                document.add_paragraph(f"[Error embedding plot: {img_err}]")
            finally:
                if os.path.exists(plot_filename): os.remove(plot_filename) 

            # Model Evaluation Metrics Section
            document.add_heading('2. Model Evaluation Metrics', level=1)
            metrics_table = document.add_table(rows=1, cols=2)
            metrics_table.style = 'Table Grid' 
            hdr_cells = metrics_table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Value"
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                
            metrics_data = {
                "RMSE": f"{st.session_state.evaluation_metrics.get('RMSE', np.nan):.4f}" if not np.isnan(st.session_state.evaluation_metrics.get('RMSE', np.nan)) else "N/A",
                "MAE": f"{st.session_state.evaluation_metrics.get('MAE', np.nan):.4f}" if not np.isnan(st.session_state.evaluation_metrics.get('MAE', np.nan)) else "N/A",
                "MAPE": f"{st.session_state.evaluation_metrics.get('MAPE', np.nan):.2f}%" if not np.isnan(st.session_state.evaluation_metrics.get('MAPE', np.nan)) else "N/A"
            }
            for metric, value in metrics_data.items():
                row_cells = metrics_table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
                for cell in row_cells: 
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            document.add_paragraph("\n")

            # Forecast Data Table Section
            document.add_heading('3. Forecast Data (First 10 rows)', level=1)
            forecast_table = document.add_table(rows=1, cols=4)
            forecast_table.style = 'Table Grid'
            hdr_cells = forecast_table.rows[0].cells
            hdr_cells[0].text = "Date"
            hdr_cells[1].text = "Forecast"
            hdr_cells[2].text = "Lower CI"
            hdr_cells[3].text = "Upper CI"
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)

            for _, row in st.session_state.forecast_results.head(10).iterrows():
                row_cells = forecast_table.add_row().cells
                row_cells[0].text = str(row["Date"].date())
                row_cells[1].text = f"{row['Forecast']:.2f}"
                row_cells[2].text = f"{row['Lower_CI']:.2f}"
                row_cells[3].text = f"{row['Upper_CI']:.2f}"
                for cell in row_cells: 
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            document.add_paragraph("\n")

            # AI Report Section
            document.add_heading(f'4. AI Report ({st.session_state.report_language})', level=1)
            for para_text in st.session_state.ai_report.split('\n\n'): 
                if para_text.strip():
                    document.add_paragraph(para_text.strip())
            document.add_paragraph("\n")
            
            # Save document to BytesIO
            document.save(docx_buffer)
            docx_buffer.seek(0) # Rewind buffer for download
            
            st.download_button(
                label="Download Report (DOCX)", 
                data=docx_buffer, 
                file_name="deephydro_forecast_report.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                key="docx_download_final_btn_tab", # Renamed key to avoid conflict if any from sidebar
                use_container_width=True
            )
            st.success("Word document generated. Click the button above to download.")
            if firebase_initialized: log_visitor_activity("Main Area", "docx_download_success")
        else: 
            st.info("Run a forecast and generate AI report first to enable DOCX download.")
    else: 
        st.info("Click 'Generate AI Report' (sidebar) after a forecast.")


# AI Chatbot Tab
with tabs[5]:
    if firebase_initialized: log_visitor_activity("Tab: AI Chatbot", "view")
    st.header("AI Chatbot Assistant")
    if not gemini_configured: st.warning("AI features disabled. Configure Gemini API Key.")
    elif st.session_state.chat_active:
        if st.session_state.cleaned_data is not None and st.session_state.forecast_results is not None and st.session_state.evaluation_metrics is not None:
            st.info("Chat activated. Ask about the results.")
            chat_container = st.container(height=400) 
            with chat_container:
                for sender, message in st.session_state.chat_history:
                    msg_class = "user-message" if sender == "User" else "ai-message"
                    st.markdown(f'<div class="chat-message {msg_class}">{message}<span class="copy-tooltip">Copied!</span></div>', unsafe_allow_html=True)
            
            user_input = st.chat_input("Ask the AI assistant:")
            if user_input:
                if firebase_initialized: log_visitor_activity("Chat", "send_message")
                st.session_state.chat_history.append(("User", user_input))
                with st.spinner("AI thinking..."): 
                    ai_response = get_gemini_chat_response(
                        user_input, st.session_state.chat_history, st.session_state.cleaned_data,
                        st.session_state.forecast_results, st.session_state.evaluation_metrics, st.session_state.ai_report
                    )
                st.session_state.chat_history.append(("AI", ai_response))
                st.rerun()
        else:
            st.warning("Run a successful forecast first to provide context for the chatbot.")
            st.session_state.chat_active = False
            st.rerun()
    else:
        st.info("Click 'Activate Chat' (sidebar) after a forecast." if gemini_configured else "AI Chat disabled.")

# Admin Analytics Tab
with tabs[6]:
    if firebase_initialized: log_visitor_activity("Tab: Admin Analytics", "view")
    render_admin_analytics()

